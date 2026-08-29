import builtins

# Compatibility for older packages
if not hasattr(builtins, "xrange"):
    builtins.xrange = range

import os
import re
from io import BytesIO

import pypdf
import docx


class ResumeParser:
    def __init__(self):
        pass

    # =========================================================
    # PDF TEXT EXTRACTION
    # =========================================================

    def extract_text_from_pdf(self, pdf_file):
        """
        Extract text from a PDF file using pypdf.

        Supports:
        - Flask FileStorage
        - BytesIO
        - raw bytes
        """

        try:
            # -------------------------------------------------
            # Convert input into bytes
            # -------------------------------------------------

            if hasattr(pdf_file, "read"):
                pdf_file.seek(0)
                file_content = pdf_file.read()
                pdf_file.seek(0)
            elif isinstance(pdf_file, bytes):
                file_content = pdf_file
            else:
                raise ValueError("Unsupported PDF input type.")

            if not file_content:
                raise ValueError("The uploaded PDF is empty.")

            # -------------------------------------------------
            # Create PDF reader
            # -------------------------------------------------

            pdf_stream = BytesIO(file_content)

            reader = pypdf.PdfReader(pdf_stream)

            # Check encrypted PDF
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception:
                    raise ValueError(
                        "This PDF is password protected. "
                        "Please upload an unlocked PDF."
                    )

            # -------------------------------------------------
            # Extract text page by page
            # -------------------------------------------------

            text_parts = []

            for page_number, page in enumerate(reader.pages, start=1):

                try:
                    page_text = page.extract_text()

                    if page_text:
                        page_text = page_text.strip()

                        if page_text:
                            text_parts.append(page_text)

                except Exception as page_error:

                    print(
                        f"Warning: Could not extract text "
                        f"from PDF page {page_number}: {page_error}"
                    )

            # -------------------------------------------------
            # Combine text
            # -------------------------------------------------

            text = "\n\n".join(text_parts)

            # Normalize excessive whitespace
            text = self.clean_text(text)

            # -------------------------------------------------
            # IMPORTANT:
            # If no text was extracted, don't silently hide it.
            # -------------------------------------------------

            if not text.strip():

                print(
                    "\n========================================"
                    "\nPDF TEXT EXTRACTION RESULT"
                    "\n========================================"
                    "\nNo text was extracted from this PDF."
                    "\nPossible reasons:"
                    "\n1. PDF is scanned/image based."
                    "\n2. PDF contains text as images."
                    "\n3. PDF uses an unusual encoding."
                    "\n4. PDF is corrupted."
                    "\n========================================\n"
                )

            return text

        except Exception as e:

            print(
                "\n========================================"
                "\nPDF EXTRACTION ERROR"
                "\n========================================"
            )

            print(str(e))

            print(
                "========================================\n"
            )

            return ""

    # =========================================================
    # DOCX TEXT EXTRACTION
    # =========================================================

    def extract_text_from_docx(self, docx_file):
        """
        Extract text from DOCX files.

        Extracts:
        - paragraphs
        - tables
        """

        try:

            # -------------------------------------------------
            # Convert input to BytesIO
            # -------------------------------------------------

            if hasattr(docx_file, "read"):

                docx_file.seek(0)

                file_content = docx_file.read()

                docx_file.seek(0)

            elif isinstance(docx_file, bytes):

                file_content = docx_file

            else:

                raise ValueError("Unsupported DOCX input type.")

            if not file_content:

                raise ValueError("The uploaded DOCX file is empty.")

            # -------------------------------------------------
            # Open DOCX
            # -------------------------------------------------

            document = docx.Document(BytesIO(file_content))

            text_parts = []

            # -------------------------------------------------
            # Extract paragraphs
            # -------------------------------------------------

            for paragraph in document.paragraphs:

                paragraph_text = paragraph.text.strip()

                if paragraph_text:

                    text_parts.append(paragraph_text)

            # -------------------------------------------------
            # Extract tables
            #
            # Many resumes contain information inside tables.
            # -------------------------------------------------

            for table in document.tables:

                for row in table.rows:

                    row_text = []

                    for cell in row.cells:

                        cell_text = cell.text.strip()

                        if cell_text:

                            row_text.append(cell_text)

                    if row_text:

                        text_parts.append(" | ".join(row_text))

            # -------------------------------------------------
            # Combine
            # -------------------------------------------------

            text = "\n".join(text_parts)

            text = self.clean_text(text)

            if not text.strip():

                print(
                    "\n========================================"
                    "\nDOCX TEXT EXTRACTION RESULT"
                    "\n========================================"
                    "\nNo text was extracted from this DOCX."
                    "\n========================================\n"
                )

            return text

        except Exception as e:

            print(
                "\n========================================"
                "\nDOCX EXTRACTION ERROR"
                "\n========================================"
            )

            print(str(e))

            print(
                "========================================\n"
            )

            return ""

    # =========================================================
    # GENERAL TEXT EXTRACTION
    # =========================================================

    def extract_text(self, file):
        """
        Detect file type and extract text.

        Used by app.py:

            resume_text = parser.extract_text(uploaded)
        """

        try:

            # -------------------------------------------------
            # Validate file
            # -------------------------------------------------

            if file is None:

                print("No file was provided.")

                return ""

            # -------------------------------------------------
            # Get filename safely
            # -------------------------------------------------

            filename = getattr(file, "filename", "")

            if not filename:

                filename = getattr(file, "name", "")

            filename = filename.lower().strip()

            print(
                f"\nResumeParser: Processing file -> {filename}"
            )

            # -------------------------------------------------
            # Reset file pointer
            # -------------------------------------------------

            if hasattr(file, "seek"):

                file.seek(0)

            # -------------------------------------------------
            # Determine extension
            # -------------------------------------------------

            extension = os.path.splitext(filename)[1].lower()

            # -------------------------------------------------
            # PDF
            # -------------------------------------------------

            if extension == ".pdf":

                text = self.extract_text_from_pdf(file)

            # -------------------------------------------------
            # DOCX
            # -------------------------------------------------

            elif extension == ".docx":

                text = self.extract_text_from_docx(file)

            # -------------------------------------------------
            # Unsupported
            # -------------------------------------------------

            else:

                print(
                    f"Unsupported file type: {extension}"
                )

                return ""

            # -------------------------------------------------
            # Final validation
            # -------------------------------------------------

            if text:

                print(
                    f"ResumeParser: Successfully extracted "
                    f"{len(text)} characters."
                )

                print(
                    "First 300 characters:"
                )

                print(
                    repr(text[:300])
                )

            else:

                print(
                    "ResumeParser: Extracted text is EMPTY."
                )

            return text

        except Exception as e:

            print(
                "\n========================================"
                "\nRESUME PARSER ERROR"
                "\n========================================"
            )

            print(str(e))

            print(
                "========================================\n"
            )

            return ""

    # =========================================================
    # CLEAN TEXT
    # =========================================================

    def clean_text(self, text):
        """
        Clean extracted text without destroying useful
        resume information.
        """

        if not text:

            return ""

        # Normalize Windows line endings
        text = text.replace("\r\n", "\n")

        # Normalize old Mac line endings
        text = text.replace("\r", "\n")

        # Replace non-breaking spaces
        text = text.replace("\xa0", " ")

        # Remove excessive spaces
        text = re.sub(r"[ \t]+", " ", text)

        # Remove excessive blank lines
        text = re.sub(r"\n{3,}", "\n\n", text)

        return text.strip()

    # =========================================================
    # SIMPLE RESUME PARSING
    # =========================================================

    def parse(self, file):

        text = self.extract_text(file)

        # -----------------------------------------------------
        # If extraction failed
        # -----------------------------------------------------

        if not text:

            return {
                "skills": [],
                "experience": [],
                "education": [],
                "raw_text": ""
            }

        # -----------------------------------------------------
        # Skills
        # -----------------------------------------------------

        skills = []

        skill_keywords = [

            # Programming
            "python",
            "java",
            "javascript",
            "typescript",
            "c",
            "c++",
            "c#",

            # Web
            "html",
            "css",
            "react",
            "angular",
            "vue",
            "next.js",
            "node.js",
            "node",
            "express",
            "flask",
            "django",
            "spring",
            "spring boot",

            # Database
            "sql",
            "mysql",
            "postgresql",
            "postgres",
            "mongodb",
            "oracle",
            "redis",

            # Data / ML
            "machine learning",
            "ml"
            "deep learning",
            "data science",
            "data analysis",
            "pandas",
            "numpy",
            "scikit-learn",
            "tensorflow",
            "pytorch",
            "keras",
            "matplotlib",
            "seaborn",

            # DevOps / Cloud
            "docker",
            "kubernetes",
            "aws",
            "azure",
            "gcp",
            "jenkins",
            "git",
            "github",
            "gitlab",

            # Tools
            "jira",
            "linux",
            "power bi",
            "tableau",
            "excel"
        ]

        text_lower = text.lower()

        for skill in skill_keywords:

            if skill.lower() in text_lower:

                skills.append(skill)

        # -----------------------------------------------------
        # Experience
        # -----------------------------------------------------

        experience = []

        lines = text.split("\n")

        experience_keywords = [
            "experience",
            "work experience",
            "professional experience",
            "employment",
            "work history",
            "internship"
        ]

        education_keywords = [
            "education",
            "academic",
            "qualification",
            "degree",
            "university",
            "college",
            "school"
        ]

        current_section = None

        for line in lines:

            clean_line = line.strip()

            if not clean_line:
                continue

            lower_line = clean_line.lower()

            # Detect experience section
            if any(
                keyword in lower_line
                for keyword in experience_keywords
            ):

                current_section = "experience"
                continue

            # Detect education section
            if any(
                keyword in lower_line
                for keyword in education_keywords
            ):

                current_section = "education"
                continue

            # Add content
            if current_section == "experience":

                experience.append(clean_line)

        # -----------------------------------------------------
        # Education
        # -----------------------------------------------------

        education = []

        current_section = None

        for line in lines:

            clean_line = line.strip()

            if not clean_line:
                continue

            lower_line = clean_line.lower()

            if any(
                keyword in lower_line
                for keyword in education_keywords
            ):

                current_section = "education"
                continue

            if any(
                keyword in lower_line
                for keyword in experience_keywords
            ):

                current_section = "experience"
                continue

            if current_section == "education":

                education.append(clean_line)

        # -----------------------------------------------------
        # Return parsed result
        # -----------------------------------------------------

        return {
            "skills": skills,
            "experience": experience,
            "education": education,
            "raw_text": text
        }